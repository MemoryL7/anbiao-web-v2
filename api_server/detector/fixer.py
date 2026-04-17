"""
fixer.py — Auto-fix logic for docx dark-bid issues.

Supported fixes:
  1. Margins          — set all section margins to config value
  2. Font name        — change non-allowed fonts to first allowed font
  3. Font size        — change non-standard font sizes to standard
  4. Line spacing     — set fixed line-spacing per config
  5. Page orientation  — force portrait
  6. Header/footer    — clear all header/footer content
  7. Page numbers     — (covered by header/footer clearing)

Notes:
  * Works on a ``docx.Document`` object in-memory.
  * Does NOT auto-delete sensitive content — only flags it.
  * Returns the modified document and a list of applied fixes.
"""

import io
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from .config import get_config


class DocxFixer:
    """Apply automatic fixes to a docx document."""

    def __init__(self, doc: Document, config: dict):
        """
        Args:
            doc: a ``docx.Document`` object (will be mutated in-place).
            config: detection config dict (from ``get_config``).
        """
        self.doc = doc
        self.config = config
        self.fixes_applied: list[str] = []

    # ── Public entry point ──────────────────────────────────────────
    def fix_all(self) -> tuple[Document, list[str]]:
        """Run all enabled fixes and return ``(doc, fixes_applied)``."""
        self.fixes_applied = []
        self.fix_margins()
        self.fix_orientation()
        self.fix_font_name()
        self.fix_font_size()
        self.fix_line_spacing()
        self.fix_headers_footers()
        return self.doc, self.fixes_applied

    # ── Margins ─────────────────────────────────────────────────────
    def fix_margins(self):
        page_cfg = self.config.get("page_format", {})
        std_cm = page_cfg.get("margin_standard")
        if std_cm is None:
            return
        for section in self.doc.sections:
            changed = False
            for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
                current = getattr(section, attr)
                target = Cm(std_cm)
                if abs(current.cm - std_cm) > 0.01:
                    setattr(section, attr, target)
                    changed = True
            if changed:
                self.fixes_applied.append(f"页边距已统一设置为 {std_cm} cm")

    # ── Orientation ─────────────────────────────────────────────────
    def fix_orientation(self):
        page_cfg = self.config.get("page_format", {})
        expected = page_cfg.get("orientation")
        if expected is None:
            return
        target_ori = WD_ORIENT.PORTRAIT if expected == "纵向" else WD_ORIENT.LANDSCAPE
        for section in self.doc.sections:
            if section.orientation != target_ori:
                # Swap width/height when switching orientation
                if target_ori == WD_ORIENT.PORTRAIT:
                    w, h = section.page_width, section.page_height
                    if w > h:
                        section.orientation = WD_ORIENT.PORTRAIT
                        section.page_width, section.page_height = h, w
                else:
                    w, h = section.page_width, section.page_height
                    if h > w:
                        section.orientation = WD_ORIENT.LANDSCAPE
                        section.page_width, section.page_height = h, w
                self.fixes_applied.append(f"页面方向已设置为 {expected}")

    # ── Font name ───────────────────────────────────────────────────
    def fix_font_name(self):
        font_cfg = self.config.get("font_format", {})
        allowed = font_cfg.get("allowed_fonts")
        if not allowed:
            return
        target_font = allowed[0]
        changed = False
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.name and run.font.name not in allowed:
                    run.font.name = target_font
                    # Also set East-Asian font via XML
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.find(qn("w:rFonts"))
                    if rfonts is None:
                        rfonts = OxmlElement("w:rFonts")
                        rpr.insert(0, rfonts)
                    rfonts.set(qn("w:eastAsia"), target_font)
                    rfonts.set(qn("w:ascii"), target_font)
                    rfonts.set(qn("w:hAnsi"), target_font)
                    changed = True
            # Also check paragraph-level font via style
            if para.style and para.style.font:
                if para.style.font.name and para.style.font.name not in allowed:
                    para.style.font.name = target_font
                    changed = True
        if changed:
            self.fixes_applied.append(f"字体已统一设置为 {target_font}")

    # ── Font size ───────────────────────────────────────────────────
    def fix_font_size(self):
        font_cfg = self.config.get("font_format", {})
        std_size = font_cfg.get("font_size")
        if std_size is None:
            return
        target = Pt(std_size)
        changed = False
        for para in self.doc.paragraphs:
            for run in para.runs:
                if run.font.size and abs(run.font.size.pt - std_size) > 0.01:
                    run.font.size = target
                    changed = True
        if changed:
            self.fixes_applied.append(f"字体大小已统一设置为 {std_size} pt")

    # ── Line spacing ────────────────────────────────────────────────
    def fix_line_spacing(self):
        font_cfg = self.config.get("font_format", {})
        ls_cfg = font_cfg.get("line_spacing", {})
        expected_type = ls_cfg.get("type")
        expected_value = ls_cfg.get("value")
        if expected_type is None:
            return

        changed = False
        for para in self.doc.paragraphs:
            pf = para.paragraph_format
            if expected_type == "固定值":
                target_rule = WD_LINE_SPACING.EXACTLY
                target_value = Pt(expected_value) if expected_value else None
            elif expected_type == "单倍行距":
                target_rule = WD_LINE_SPACING.SINGLE
                target_value = None
            elif expected_type == "1.5倍行距":
                target_rule = WD_LINE_SPACING.ONE_POINT_FIVE
                target_value = None
            elif expected_type == "2倍行距":
                target_rule = WD_LINE_SPACING.DOUBLE
                target_value = None
            else:
                # Default to fixed
                target_rule = WD_LINE_SPACING.EXACTLY
                target_value = Pt(expected_value) if expected_value else None

            if pf.line_spacing_rule != target_rule:
                pf.line_spacing_rule = target_rule
                changed = True
            if target_value is not None:
                if pf.line_spacing != expected_value:
                    pf.line_spacing = expected_value
                    changed = True
        if changed:
            desc = f"{expected_type}"
            if expected_value:
                desc += f" {expected_value} pt"
            self.fixes_applied.append(f"行间距已设置为 {desc}")

    # ── Headers / Footers ──────────────────────────────────────────
    def fix_headers_footers(self):
        page_cfg = self.config.get("page_format", {})
        if not page_cfg.get("forbid_header_footer", True):
            return
        changed = False
        for section in self.doc.sections:
            for hf in (section.header, section.footer):
                for para in hf.paragraphs:
                    if para.text.strip():
                        para.clear()
                        changed = True
                # Also clear tables in header/footer
                for table in hf.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                if para.text.strip():
                                    para.clear()
                                    changed = True
            # Remove the "link to previous" if it just copies content
            # (This is a pass — headers/footers are already cleared)
        if changed:
            self.fixes_applied.append("页眉页脚内容已清除")

    # ── Utilities ───────────────────────────────────────────────────
    def save_to_bytes(self) -> bytes:
        """Save the fixed document to bytes."""
        buf = io.BytesIO()
        self.doc.save(buf)
        return buf.getvalue()
