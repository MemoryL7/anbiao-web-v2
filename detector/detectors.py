"""
检测器模块 - PageFormatDetector, FontFormatDetector, TitleFormatDetector, ContentChecker
Ported from utils/*.py with identical detection logic.
"""

import re
from docx.oxml.ns import qn


# ╔═══════════════════════════════════════════════════════════════════╗
# ║  PageFormatDetector                                              ║
# ╚═══════════════════════════════════════════════════════════════════╝
class PageFormatDetector:
    """Detect page-format issues: margins, orientation, headers/footers."""

    def __init__(self, doc_object, config: dict):
        self.doc = doc_object
        self.config = config
        self.results: list[dict] = []

    def check_page_format(self) -> list[dict]:
        self.results = []
        self.check_margins()
        self.check_header_footer()
        self.check_orientation()
        return self.results

    # ── Margins ─────────────────────────────────────────────────────
    def check_margins(self):
        if not self.doc or not self.config:
            return
        page_cfg = self.config.get("page_format", {})
        std = page_cfg.get("margin_standard", 2.5)
        tol = page_cfg.get("margin_tolerance", 0.0)
        std_inch = std / 2.54
        tol_inch = tol / 2.54

        for i, section in enumerate(self.doc.sections):
            margins = {
                "top": section.top_margin.inches,
                "bottom": section.bottom_margin.inches,
                "left": section.left_margin.inches,
                "right": section.right_margin.inches,
            }
            name_map = {"top": "上", "bottom": "下", "left": "左", "right": "右"}
            lo, hi = std_inch - tol_inch, std_inch + tol_inch
            for mtype, mval in margins.items():
                if not (lo <= mval <= hi):
                    self.results.append(
                        {
                            "type": "page_format",
                            "subtype": "margin",
                            "section": i + 1,
                            "margin_type": name_map[mtype],
                            "actual": round(mval * 2.54, 2),
                            "expected": std,
                            "message": (
                                f"第{i+1}节{name_map[mtype]}边距不符合要求，"
                                f"应为{std}cm，实际为{round(mval * 2.54, 2)}cm"
                            ),
                        }
                    )

    # ── Orientation ─────────────────────────────────────────────────
    def check_orientation(self):
        if not self.doc:
            return
        for i, section in enumerate(self.doc.sections):
            ori = section.orientation
            if ori != 0:  # 0=PORTRAIT
                self.results.append(
                    {
                        "type": "page_format",
                        "subtype": "orientation",
                        "section": i + 1,
                        "actual": "横向",
                        "expected": "纵向",
                        "message": f"第{i+1}节页面方向不符合要求，应为纵向，实际为横向",
                    }
                )

    # ── Header / Footer ────────────────────────────────────────────
    def check_header_footer(self):
        if not self.doc or not self.config:
            return
        page_cfg = self.config.get("page_format", {})
        if page_cfg.get("forbid_header_footer", True):
            for i, section in enumerate(self.doc.sections):
                header = section.header
                if header.paragraphs and any(p.text.strip() for p in header.paragraphs):
                    self.results.append(
                        {
                            "type": "page_format",
                            "subtype": "header",
                            "section": i + 1,
                            "message": f"第{i+1}节包含页眉，不符合要求",
                        }
                    )
                footer = section.footer
                if footer.paragraphs and any(p.text.strip() for p in footer.paragraphs):
                    self.results.append(
                        {
                            "type": "page_format",
                            "subtype": "footer",
                            "section": i + 1,
                            "message": f"第{i+1}节包含页脚，不符合要求",
                        }
                    )

        if page_cfg.get("forbid_page_numbers", True):
            for i, section in enumerate(self.doc.sections):
                for hf in [section.header, section.footer]:
                    for para in hf.paragraphs:
                        if "PAGE" in para._element.xml or "页" in para.text:
                            self.results.append(
                                {
                                    "type": "page_format",
                                    "subtype": "page_number",
                                    "section": i + 1,
                                    "message": f"第{i+1}节包含页码，不符合要求",
                                }
                            )
                            break

    # ── Blank pages ─────────────────────────────────────────────────
    def check_blank_pages(self):
        if not self.doc or not self.config:
            return
        page_cfg = self.config.get("page_format", {})
        if not page_cfg.get("forbid_blank_pages", True):
            return
        page_breaks = 1
        current: list = []
        for para in self.doc.paragraphs:
            xml = para._element.xml
            if "w:br" in xml and 'type="page"' in xml:
                if not any(p.text.strip() for p in current):
                    self.results.append(
                        {
                            "type": "page_format",
                            "subtype": "blank_page",
                            "page": page_breaks,
                            "message": f"第{page_breaks}页为空白页，不符合要求",
                        }
                    )
                current = []
                page_breaks += 1
            else:
                current.append(para)
        if not any(p.text.strip() for p in current):
            self.results.append(
                {
                    "type": "page_format",
                    "subtype": "blank_page",
                    "page": page_breaks,
                    "message": f"第{page_breaks}页为空白页，不符合要求",
                }
            )


# ╔═══════════════════════════════════════════════════════════════════╗
# ║  FontFormatDetector                                              ║
# ╚═══════════════════════════════════════════════════════════════════╝
class FontFormatDetector:
    """Detect font-format issues: font name, size, line spacing, italic, underline."""

    def __init__(self, doc_object, config: dict):
        self.doc = doc_object
        self.config = config
        self.results: list[dict] = []

    def check_font_format(self) -> list[dict]:
        self.results = []
        self.check_font_name()
        self.check_font_size()
        self.check_line_spacing()
        return self.results

    # ── Font name ───────────────────────────────────────────────────
    def check_font_name(self):
        if not self.doc or not self.config:
            return
        allowed = self.config.get("font_format", {}).get("allowed_fonts", ["宋体"])
        for pi, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            for ri, run in enumerate(para.runs):
                if not run.text.strip():
                    continue
                fname = run.font.name
                if fname and fname not in allowed:
                    self.results.append(
                        {
                            "type": "font_format",
                            "subtype": "font_name",
                            "paragraph": pi + 1,
                            "run": ri + 1,
                            "actual": fname,
                            "expected": allowed,
                            "text": run.text.strip(),
                            "message": (
                                f"第{pi+1}段第{ri+1}处使用了不允许的字体：{fname}，"
                                f"允许字体：{allowed}"
                            ),
                        }
                    )

    # ── Font size ───────────────────────────────────────────────────
    def check_font_size(self):
        if not self.doc or not self.config:
            return
        std_size = self.config.get("font_format", {}).get("font_size", 14.0)
        for pi, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            for ri, run in enumerate(para.runs):
                if not run.text.strip():
                    continue
                fs = run.font.size
                if fs and fs.pt != std_size:
                    self.results.append(
                        {
                            "type": "font_format",
                            "subtype": "font_size",
                            "paragraph": pi + 1,
                            "run": ri + 1,
                            "actual": fs.pt,
                            "expected": std_size,
                            "text": run.text.strip(),
                            "message": (
                                f"第{pi+1}段第{ri+1}处字体大小不符合要求：{fs.pt}pt，"
                                f"应为：{std_size}pt"
                            ),
                        }
                    )

    # ── Line spacing ────────────────────────────────────────────────
    def check_line_spacing(self):
        if not self.doc or not self.config:
            return
        ls_cfg = self.config.get("font_format", {}).get("line_spacing", {})
        expected_type = ls_cfg.get("type", "固定值")
        expected_value = ls_cfg.get("value", 28.0)

        rule_map = {0: "单倍行距", 1: "1.5倍行距", 2: "2倍行距", 3: "多倍行距", 4: "固定值", 5: "最小值"}

        for pi, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue
            pf = para.paragraph_format
            actual_type = rule_map.get(pf.line_spacing_rule, "未知")
            actual_value = pf.line_spacing

            if expected_type == "固定值":
                if actual_type != expected_type:
                    self.results.append(
                        {
                            "type": "font_format",
                            "subtype": "line_spacing_type",
                            "paragraph": pi + 1,
                            "actual": actual_type,
                            "expected": expected_type,
                            "message": f"第{pi+1}段行间距类型不符合要求：{actual_type}，应为：{expected_type}",
                        }
                    )
                elif actual_value != expected_value:
                    self.results.append(
                        {
                            "type": "font_format",
                            "subtype": "line_spacing_value",
                            "paragraph": pi + 1,
                            "actual": actual_value,
                            "expected": expected_value,
                            "message": f"第{pi+1}段行间距值不符合要求：{actual_value}pt，应为：{expected_value}pt",
                        }
                    )

            # Italic check
            if self.config.get("font_format", {}).get("forbid_italic", True):
                for run in para.runs:
                    if run.italic:
                        self.results.append(
                            {
                                "type": "font_format",
                                "subtype": "italic",
                                "paragraph": pi + 1,
                                "text": run.text.strip(),
                                "message": f"第{pi+1}段包含不允许的斜体字体",
                            }
                        )

            # Underline check
            if self.config.get("font_format", {}).get("forbid_underline", True):
                for run in para.runs:
                    if run.underline:
                        self.results.append(
                            {
                                "type": "font_format",
                                "subtype": "underline",
                                "paragraph": pi + 1,
                                "text": run.text.strip(),
                                "message": f"第{pi+1}段包含不允许的下划线字体",
                            }
                        )


# ╔═══════════════════════════════════════════════════════════════════╗
# ║  TitleFormatDetector                                             ║
# ╚═══════════════════════════════════════════════════════════════════╝
class TitleFormatDetector:
    """Detect title-format issues: numbering patterns and hierarchy."""

    def __init__(self, doc_object, config: dict):
        self.doc = doc_object
        self.config = config
        self.results: list[dict] = []

    def check_title_format(self) -> list[dict]:
        self.results = []
        self.check_title_patterns()
        self.check_title_hierarchy()
        return self.results

    # ── Patterns ────────────────────────────────────────────────────
    def check_title_patterns(self):
        if not self.doc or not self.config:
            return
        title_cfg = self.config.get("title_format", {})
        patterns = title_cfg.get("patterns", [r"^[一二三四五六七八九十]+、", r"^\([一二三四五六七八九十]+\)"])

        for pi, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            is_title = any(r.font.size and r.font.size.pt > 14 for r in para.runs)
            if not is_title:
                continue
            if not any(re.match(p, text) for p in patterns):
                self.results.append(
                    {
                        "type": "title_format",
                        "subtype": "pattern",
                        "paragraph": pi + 1,
                        "title_text": text,
                        "expected_patterns": patterns,
                        "message": f'第{pi+1}段标题格式不符合要求："{text}"，应匹配模式：{patterns}',
                    }
                )

    # ── Hierarchy ───────────────────────────────────────────────────
    def check_title_hierarchy(self):
        if not self.doc or not self.config:
            return
        title_cfg = self.config.get("title_format", {})
        max_levels = title_cfg.get("max_levels", 7)
        patterns = title_cfg.get("patterns", [])

        titles = []
        for pi, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            is_title = any(r.font.size and r.font.size.pt > 14 for r in para.runs)
            if not is_title:
                continue
            level = None
            for li, pat in enumerate(patterns):
                if re.match(pat, text):
                    level = li + 1
                    break
            if level:
                titles.append({"paragraph": pi + 1, "text": text, "level": level})

        if len(titles) > 1:
            for i in range(1, len(titles)):
                cur, prev = titles[i], titles[i - 1]
                if cur["level"] > prev["level"] + 1:
                    self.results.append(
                        {
                            "type": "title_format",
                            "subtype": "hierarchy",
                            "current_paragraph": cur["paragraph"],
                            "current_title": cur["text"],
                            "current_level": cur["level"],
                            "previous_paragraph": prev["paragraph"],
                            "previous_title": prev["text"],
                            "previous_level": prev["level"],
                            "message": (
                                f'标题层级不合理：第{cur["paragraph"]}段标题级别为{cur["level"]}，'
                                f'与前一个标题（第{prev["paragraph"]}段，级别{prev["level"]}）相差超过1级'
                            ),
                        }
                    )
                if cur["level"] > max_levels:
                    self.results.append(
                        {
                            "type": "title_format",
                            "subtype": "max_level",
                            "paragraph": cur["paragraph"],
                            "title_text": cur["text"],
                            "actual_level": cur["level"],
                            "max_levels": max_levels,
                            "message": f'标题级别超过最大允许级别：第{cur["paragraph"]}段标题级别为{cur["level"]}，最大允许级别为{max_levels}',
                        }
                    )


# ╔═══════════════════════════════════════════════════════════════════╗
# ║  ContentChecker                                                  ║
# ╚═══════════════════════════════════════════════════════════════════╝
class ContentChecker:
    """Detect content-compliance issues: sensitive words, images, signatures."""

    def __init__(
        self,
        doc_object,
        config: dict,
        check_images: bool = True,
        custom_sensitive_words: list[str] | None = None,
    ):
        self.doc = doc_object
        self.config = config
        self.check_images_flag = check_images
        self.custom_sensitive_words = custom_sensitive_words or []
        self.results: list[dict] = []

    def check_content(self) -> list[dict]:
        self.results = []
        self.check_sensitive_info()
        if self.check_images_flag:
            self.check_images()
        self.check_electronic_signature()
        return self.results

    # ── Sensitive info (flag only, never delete) ────────────────────
    def check_sensitive_info(self):
        if not self.doc or not self.config:
            return
        patterns = self.config.get("content_check", {}).get("sensitive_patterns", [])
        all_patterns = patterns + self.custom_sensitive_words
        if not all_patterns:
            return
        for pi, para in enumerate(self.doc.paragraphs):
            text = para.text
            if not text:
                continue
            for pat in all_patterns:
                for m in re.finditer(pat, text, re.IGNORECASE):
                    self.results.append(
                        {
                            "type": "content_check",
                            "subtype": "sensitive_info",
                            "paragraph": pi + 1,
                            "sensitive_text": m.group(),
                            "pattern": pat,
                            "message": f'第{pi+1}段包含敏感信息："{m.group()}"，匹配模式：{pat}',
                        }
                    )

    # ── Images ──────────────────────────────────────────────────────
    def check_images(self):
        if not self.doc or not self.config:
            return
        if not self.config.get("content_check", {}).get("forbid_images", True):
            return
        for pi, para in enumerate(self.doc.paragraphs):
            if "Graphic" in para._element.xml:
                self.results.append(
                    {
                        "type": "content_check",
                        "subtype": "image",
                        "paragraph": pi + 1,
                        "message": f"第{pi+1}段包含图片，不符合要求",
                    }
                )
        try:
            for ii, shape in enumerate(self.doc.inline_shapes):
                if shape.type == 3:
                    self.results.append(
                        {
                            "type": "content_check",
                            "subtype": "image",
                            "image_index": ii + 1,
                            "message": f"文档包含嵌入式图片，索引：{ii+1}，不符合要求",
                        }
                    )
        except Exception:
            pass

    # ── Electronic signatures ───────────────────────────────────────
    def check_electronic_signature(self):
        if not self.doc or not self.config:
            return
        if not self.config.get("content_check", {}).get("forbid_electronic_signature", True):
            return
        try:
            doc_xml = self.doc._element.xml
            for tag in ("w:signature", "ds:Signature", "xades:Signature"):
                if tag in doc_xml:
                    self.results.append(
                        {
                            "type": "content_check",
                            "subtype": "electronic_signature",
                            "message": "文档包含电子签章，不符合要求",
                        }
                    )
                    break
        except Exception:
            pass
